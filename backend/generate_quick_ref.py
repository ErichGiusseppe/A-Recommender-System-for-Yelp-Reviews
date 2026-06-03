"""
generate_quick_ref.py — referencia rápida del recomendador Lantern.
Corre desde backend/: python generate_quick_ref.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ── helpers ──────────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1C, 0x19, 0x17)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    return p

def bwc(doc, parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for text, is_code in parts:
        run = p.add_run(text)
        if is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    return p

def code_block(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x57, 0x53, 0x4E)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5F5F4")
    p._p.get_or_add_pPr().append(shd)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FFF7ED")
    p._p.get_or_add_pPr().append(shd)
    return p

def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 95)
    run.font.color.rgb = RGBColor(0xE7, 0xE5, 0xE4)
    run.font.size = Pt(7)

def simple_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # header row
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1C1917")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    # data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ci == 0:
                cells[ci].paragraphs[0].runs[0].font.name = "Courier New"
                cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
            if ri % 2 == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FAF6F0")
                cells[ci]._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return tbl


# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lantern — Quick Reference: Recomendador")
run.bold = True; run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Caminos del scorer · Fórmula · Vistas → Endpoints → Funciones")
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "0. Glosario rápido — leer antes de continuar")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "SVD++ (Singular Value Decomposition++): algoritmo de filtrado colaborativo entrenado "
    "con las reseñas del dataset Yelp. Aprende un vector de 25 números ('vector latente') "
    "por cada usuario y por cada negocio. El producto punto de ambos vectores predice cuánto "
    "le va a gustar ese negocio a ese usuario. Es el modelo principal de Lantern.")

body(doc,
    "Trainset: el conjunto de datos con el que se entrenó SVD++. Solo están en el trainset "
    "los usuarios que ya existían en el dataset Yelp al momento del entrenamiento. Los 7 "
    "usuarios demo de Lantern (camila, daniel, sara, alex, maria, carlos, sofia) fueron "
    "seleccionados del trainset precisamente para que SVD++ ya tenga vectores para ellos "
    "y pueda generar recomendaciones personalizadas desde el primer login.")

body(doc,
    "Usuario warm / usuario cold: un usuario warm tiene vector en el modelo SVD++ (está en "
    "el trainset). Un usuario cold no tiene vector — es alguien que se registró en la app "
    "con un UUID nuevo que el modelo nunca vio. Para los usuarios cold, el sistema usa "
    "el wizard de preferencias en su lugar.")

body(doc,
    "Wizard de preferencias (cold-start wizard): formulario de 4 pasos que aparece al "
    "abrir la app por primera vez o al registrarse. El usuario elige: categorías de interés "
    "(Italian, Seafood, Coffee...), la ocasión habitual (viajando, salida local, cita, "
    "algo rápido), su franja horaria preferida y el rango de precio. El perfil resultante "
    "se guarda en el navegador y en el backend, y alimenta el modelo TF-IDF para generar "
    "recomendaciones de contenido cuando no hay historial colaborativo.")

body(doc,
    "Folding-in: técnica para actualizar el vector latente de un usuario sin reentrenar "
    "el modelo completo. Cuando el usuario califica negocios en la app, el sistema resuelve "
    "un sistema de mínimos cuadrados usando esas nuevas calificaciones y la matriz de "
    "factores de los negocios (que permanece fija). El resultado es un vector p_u "
    "actualizado que refleja los nuevos gustos del usuario.")

body(doc,
    "Blend progresivo (α): mecanismo de transición entre cold-start y folding-in. "
    "Con pocas reseñas el sistema no confía completamente en el vector calculado "
    "(hay pocas ecuaciones para 25 incógnitas), así que mezcla el resultado del "
    "folding-in con la señal del wizard usando α = n_reseñas / 5. Con 5 reseñas "
    "α = 1 y el wizard se ignora por completo.")

body(doc,
    "ExplanationCard: componente del detalle de negocio que muestra al usuario por qué se "
    "le recomendó ese lugar. Desglosa el score en cuatro señales: CF (cuánto vino del "
    "historial colaborativo), CB (cuánto del wizard), CTX (cuánto del contexto horario) "
    "y POP (cuánto de popularidad). El texto es generado dinámicamente según cuál señal "
    "dominó el score para ese par usuario-negocio.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. La fórmula — siempre la misma, lo que cambia es el CF")
# ══════════════════════════════════════════════════════════════════════════════

code_block(doc, "match  =  ⌊100 · (0.60 · CF  +  0.25 · CTX  +  0.15 · POP)⌋")

body(doc,
    "Los pesos 0.60 / 0.25 / 0.15 son fijos y están definidos en app/config.py "
    "(W_CF, W_CTX, W_POP). No se aprendieron automáticamente — se eligieron para dar "
    "prioridad a la señal colaborativa (60%) que es la más personalizada, con un ajuste "
    "contextual significativo (25%) y un prior de popularidad como estabilizador (15%).")

body(doc,
    "CF es la señal colaborativa: score de SVD++, folding-in o TF-IDF según el camino. "
    "Siempre va en el slot del 60% independientemente de su origen.")

body(doc,
    "CTX (contextual) es un boost predefinido por categoría de negocio y franja horaria. "
    "Los valores están en una tabla en contextual_scorer.py: por ejemplo, a las 8am "
    "un café ('coffee-and-tea') recibe CTX=90, un steakhouse recibe CTX=0. A las 10pm, "
    "los bares reciben CTX=80 y los cafés CTX=0. El CTX no se aprende de datos — "
    "refleja sentido común sobre cuándo van bien ciertos negocios. "
    "Las franjas y sus categorías representativas son:")
simple_table(doc,
    ["Franja", "Horas", "Categorías con boost alto"],
    [
        ("Morning",    "6–11h",  "coffee-and-tea (90), cafes (80), bakeries (75), donuts (70)"),
        ("Lunch",      "11–15h", "sandwiches (80), food-trucks (75), tacos (75), salad (72)"),
        ("Afternoon",  "15–18h", "coffee-and-tea (80), desserts (75), bubble-tea (72)"),
        ("Dinner",     "18–23h", "steakhouses (85), seafood (80), sushi-bars (78), barbeque (78)"),
        ("Late night", "23–6h",  "pizza (85), bars (80), diners (78), pubs (75)"),
    ]
)
body(doc,
    "Negocios sin categorías en la tabla reciben CTX=0 para esa franja. El CTX se "
    "calcula en vivo en cada request con live_ctx_score(tags, hour) — no está pre-computado.")

body(doc,
    "POP es log(1 + review_count) normalizado por el máximo de la ciudad. El logaritmo "
    "aplana la distribución: un negocio con 1000 reseñas no es 10× más popular que uno "
    "con 100. Sirve como regularizador cuando CF es débil — evita que un negocio con "
    "3 reseñas perfectas flote al tope solo por suerte estadística.")

body(doc,
    "Cuando el blend está activo (1-4 reseñas en la app), el slot del 60% se divide:")
code_block(doc, "match  =  ⌊100 · (0.60·(α·CF_folding + (1-α)·CB_wizard)  +  0.25·CTX  +  0.15·POP)⌋")
body(doc,
    "Los campos cf y cb que ve la ExplanationCard ya tienen el alpha aplicado. "
    "Con 2 reseñas (α=0.40): si CF crudo es 0.85 y CB crudo es 0.70, "
    "el usuario ve CF=34 y CB=42, no los valores crudos. Esto refleja la contribución "
    "real de cada señal al score final, no el potencial máximo de cada una.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Los 5 caminos del scorer")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "El scorer evalúa al usuario y toma el primer camino que aplique. "
    "Todo ocurre en hybrid_scorer.py → score_businesses_for_user().")

doc.add_paragraph()

# ── Camino 1 ──
h2(doc, "Camino 1 — Folding-in puro")
bwc(doc, [("¿Cuándo?", False),
    (" Usuario con ≥ 5 reseñas en la app (SQLite). Aplica igual a demos y nuevos registrados.", False)])
bwc(doc, [("Función: ", False), ("compute_folding_in_vector()", True),
    (" en svdpp_engine.py", False)])
body(doc,
    "Resuelve argmin ‖Q·p_u − r‖² con las reseñas del usuario como ecuaciones. "
    "El vector p_u nuevo se usa para puntuar todos los negocios no calificados. "
    "CF en la ExplanationCard = score SVD++ recalculado, CB = 0.")

# ── Camino 1b ──
h2(doc, "Camino 1b — Blend progresivo α")
bwc(doc, [("¿Cuándo?", False),
    (" Usuario con 1-4 reseñas en la app Y tiene perfil del wizard guardado.", False)])
bwc(doc, [("Función: ", False), ("compute_folding_in_vector()", True),
    (" + blend α aplicado en el loop de ", False), ("score_businesses_for_user()", True)])
body(doc, "α = n_reseñas / 5  →  CF_efectivo = α·folding,  CB_efectivo = (1-α)·TF-IDF")

simple_table(doc,
    ["Reseñas", "α", "CF en ExplanationCard", "CB en ExplanationCard"],
    [
        ("1", "0.20", "20% del score folding-in", "80% del score TF-IDF"),
        ("2", "0.40", "40% del score folding-in", "60% del score TF-IDF"),
        ("3", "0.60", "60% del score folding-in", "40% del score TF-IDF"),
        ("4", "0.80", "80% del score folding-in", "20% del score TF-IDF"),
        ("≥5", "1.00", "100% folding-in puro",    "0 — wizard ignorado"),
    ]
)

# ── Camino 2 ──
h2(doc, "Camino 2 — SVD++ warm (histórico)")
bwc(doc, [("¿Cuándo?", False),
    (" Usuario en el trainset SVD++ con 0 reseñas nuevas en la app (los 7 demos).", False)])
bwc(doc, [("Función: ", False), ("build_warm_user_vector()", True),
    (" + ", False), ("score_all_businesses_for_user()", True), (" en svdpp_engine.py", False)])
body(doc,
    "Recupera u_eff = p_u + (1/√|N(u)|)·Σy_j del modelo entrenado y hace Q @ u_eff en ~2ms. "
    "CF = score SVD++ histórico, CB = 0.")

# ── Camino 3 ──
h2(doc, "Camino 3 — Cold-start TF-IDF")
bwc(doc, [("¿Cuándo?", False),
    (" Usuario sin historial colaborativo: nuevo registrado (UUID no en trainset) "
     "o demo con 0 reseñas sin camino 2.", False)])
bwc(doc, [("Función: ", False), ("get_content_scores_for_city()", True),
    (" en coldstart.py → scores pre-calculados en businesses.py antes de llamar a inject_scores()", False)])
body(doc,
    "El perfil del wizard se convierte en un query TF-IDF. "
    "Similitud coseno contra vector de categorías de cada negocio. CF = 0, CB = score TF-IDF.")

# ── Camino 4 ──
h2(doc, "Camino 4 — Fallback popularidad")
bwc(doc, [("¿Cuándo?", False),
    (" Invitado sin perfil del wizard, o negocio sin cobertura en ninguna otra rama.", False)])
bwc(doc, [("Fuente: ", False), ("state.top_n['new_visitor|{city}']", True),
    (" del parquet top_n.parquet. CF y POP vienen del parquet, CTX se recalcula en vivo.", False)])

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Vistas — qué hace cada sección y cómo")
# ══════════════════════════════════════════════════════════════════════════════

note(doc,
    "Las cinco secciones de Discovery comparten los mismos 50 negocios de UNA sola llamada "
    "GET /businesses. El backend los devuelve ya ordenados por match score personal. "
    "El frontend los filtra y corta de formas distintas para cada sección — sin llamadas adicionales.")

# ── Top Picks ──
h2(doc, "Discovery — Top Picks")
bwc(doc, [("Endpoint: ", False), ("GET /businesses", True), (" → ", False),
    ("score_businesses_for_user()", True), (" [hybrid_scorer.py]", False)])
body(doc,
    "Muestra los 4 negocios con mayor match score de la lista. El backend ya los devuelve "
    "ordenados, así que el frontend simplemente toma los primeros 4.")
body(doc,
    "Excepción: si el usuario no tiene señal colaborativa real (ningún negocio de la lista "
    "tiene CF > 10, es decir es cold-start puro), el frontend hace una segunda llamada a "
    "GET /recommendations/cold-start con los parámetros del wizard. En ese caso Top Picks "
    "muestra los resultados del modelo TF-IDF en lugar del ranking SVD++. Esto evita que "
    "un usuario nuevo vea solo los negocios más populares de la ciudad sin personalización.")

sep(doc)

# ── Because You Liked ──
h2(doc, "Discovery — Because You Liked [Negocio]")
bwc(doc, [("Sin endpoint adicional — opera sobre los 50 negocios ya cargados.", False)])
body(doc,
    "El objetivo es mostrar 'más de lo que ya te gustó'. Para eso, el frontend necesita "
    "identificar QUÉ tipo de lugar le gusta más al usuario según SVD++, y luego mostrar "
    "otros del mismo tipo.")
body(doc,
    "El proceso paso a paso:")
bwc(doc, [("1. Filtrar por señal fuerte: ", False),
    ("biz.cf > 50", True),
    (" — solo negocios donde SVD++ tiene alta confianza de que le van a gustar al usuario.", False)])
body(doc,
    "2. Descartar categorías genéricas: se eliminan 'Restaurants', 'Food', 'Nightlife', 'Bars' "
    "porque no dicen nada sobre el gusto específico. Lo que queda son categorías concretas "
    "como 'Mediterranean', 'Coffee & Tea', 'Ethiopian', 'Japanese', etc.")
body(doc,
    "3. Agrupar por categoría primaria y contar: si hay 5 negocios de Mediterranean con CF>50 "
    "y 3 de Coffee & Tea, Mediterranean tiene más representación.")
body(doc,
    "4. Seleccionar la categoría dominante con mínimo de pool: la categoría elegida debe tener "
    "al menos 3 negocios con CF > 20 en la lista (anchor + 2 recomendaciones mínimo). Si solo "
    "hay 1 o 2 de ese tipo, se descarta aunque sea la más representada.")
body(doc,
    "5. El anchor: dentro de la categoría elegida, el negocio con mayor CF se convierte en "
    "el titular — 'Because you liked Zahav'. Este es el negocio que mejor representa el "
    "gusto del usuario en esa categoría según SVD++.")
body(doc,
    "6. Las recomendaciones: todos los demás negocios de esa misma categoría con CF > 20, "
    "excluyendo el anchor, ordenados por CF. Se muestran hasta 3.")
note(doc,
    "La selección entre las top 3 categorías candidatas es aleatoria en cada carga de página "
    "(estabilizada por useMemo — cambia solo cuando cambia la lista de negocios, no con cada "
    "render). Esto da variedad: si Camila tiene señal fuerte en Mediterranean Y Coffee & Tea, "
    "no siempre ve la misma. La aleatoriedad está acotada: solo elige entre candidatas con ≥3 negocios.")

sep(doc)

# ── Trending ──
h2(doc, "Discovery — Trending Nearby")
bwc(doc, [("Sin endpoint adicional.", False)])
body(doc,
    "Muestra los 3 primeros negocios de la lista ya ordenada por match, pero con un filtro "
    "de calidad mínima: rating ≥ 3.5. Sin ese filtro, un negocio con 2.5 estrellas pero "
    "muy popular podría aparecer como 'trending' si tiene CTX alto a esa hora del día. "
    "El filtro garantiza que 'trending' implique al menos calidad aceptable.")
body(doc,
    "Como la lista viene ordenada por match personal (SVD++ + CTX + POP), Trending refleja "
    "los lugares más relevantes AHORA para este usuario en esta ciudad, no simplemente "
    "los más visitados globalmente.")

sep(doc)

# ── Hidden Gems ──
h2(doc, "Discovery — Hidden Gems Nearby")
bwc(doc, [("Sin endpoint adicional.", False)])
body(doc,
    "Busca lugares excelentes pero poco conocidos. El criterio es doble: rating ≥ 4.5 "
    "(calidad real comprobada) y fewest reviews (poca visibilidad). Se ordena ascendente "
    "por número de reseñas y se toman los 3 primeros con 4.5+.")
body(doc,
    "No usa personalización: el resultado es el mismo para Camila y para Daniel en el mismo "
    "barrio. Un taller de llanta con 5 estrellas y 8 reseñas puede aparecer junto a un "
    "restaurante con 4.5 y 12 reseñas — la sección no filtra por categoría, el objetivo "
    "es descubrir cualquier negocio excelente y poco conocido.")

sep(doc)

# ── Best Value ──
h2(doc, "Discovery — Best Value")
bwc(doc, [("Sin endpoint adicional.", False)])
body(doc,
    "Muestra los 3 mejores negocios entre los económicos confirmados. El filtro es: "
    "price_known = true (el dato de precio viene del dataset Yelp real, no es un default) "
    "Y price ∈ {'$', '$$'}. Se ordena descendente por rating.")
body(doc,
    "El campo price_known existe porque ~30% de los negocios de Yelp no tienen dato de precio. "
    "Sin ese filtro, esos negocios reciben '$$' como etiqueta por defecto y aparecerían "
    "engañosamente en Best Value. Con el filtro, solo aparecen negocios donde Yelp "
    "confirma que son económicos.")

sep(doc)

# ── Search ──
h2(doc, "Search")
bwc(doc, [("Endpoint: ", False), ("GET /search", True), (" → ", False),
    ("search_businesses()", True), (" + ", False), ("inject_scores()", True)])
body(doc,
    "El usuario puede buscar por texto libre (q), filtrar por una o varias categorías "
    "simultáneamente y por precio. El backend hace dos pasos en orden:")
body(doc,
    "Primero, business_store.search_businesses() filtra el catálogo completo: "
    "busca la query en nombre, categoría, barrio y etiquetas del negocio (todo lowercase). "
    "Luego aplica los filtros de categoría y precio. Si hay filtro de ciudad, lo aplica también.")
body(doc,
    "Segundo, inject_scores() puntúa los resultados filtrados con el mismo scorer de Discovery "
    "(mismo camino SVD++/cold-start/fallback), pero SIN pasar user_ratings ni hora. "
    "Los resultados aparecen ordenados por match personal — si Camila busca 'pizza', "
    "ve las pizzerías ordenadas según su historial SVD++, no solo por popularidad.")

sep(doc)

# ── Detail ──
h2(doc, "Detail — Detalle de negocio y ExplanationCard")
bwc(doc, [("Endpoints: ", False), ("GET /businesses/{id}", True),
    (" + ", False), ("GET /explanations/{id}", True)])
body(doc,
    "Al abrir el detalle de un negocio, el frontend hace dos llamadas en paralelo.")
body(doc,
    "La primera (GET /businesses/{id}) retorna el negocio con sus scores calculados en "
    "tiempo real por inject_scores(). Esto garantiza que los valores CF/CB/CTX/POP del "
    "detail sean consistentes con los que el usuario vio en la lista.")
body(doc,
    "La segunda (GET /explanations/{id}) busca en el diccionario _explanations (cargado "
    "al startup desde explanations.parquet) por la clave 'user_id|business_id'. "
    "Es O(1). Si no hay explicación personal, hace fallback a 'new_visitor|business_id' "
    "y luego a 'camila|business_id'. Si tampoco, usa los scores del primer request.")
body(doc,
    "La ExplanationCard muestra el desglose en porcentajes y un texto generado en "
    "_apply_hybrid_score() según la señal dominante: si CF domina → 'mostly your taste "
    "history'; si CTX domina → 'perfect for this time of day'. El texto whyPicked de "
    "las cards de lista también viene de ahí.")

sep(doc)

# ── Profile ──
h2(doc, "Profile (You) — Tabs Saved y Reviews")
bwc(doc, [("Endpoints: ", False), ("GET /users/me", True),
    (" (Saved) + ", False), ("GET /reviews/me", True), (" (Reviews)", False)])
body(doc,
    "El tab Saved llama a GET /users/me que retorna entre otros datos los primeros 12 "
    "business_ids de las recomendaciones pre-computadas del parquet (get_recommendations "
    "del módulo recommender). El frontend fetcha cada uno con GET /businesses/{id} "
    "para obtener imagen y datos visuales.")
body(doc,
    "El tab Reviews llama a GET /reviews/me que combina dos fuentes: primero las reseñas "
    "escritas en la app (SQLite, prioridad), luego el historial Yelp de los 7 usuarios "
    "demo (demo_user_reviews.parquet — Camila tiene 3,048, Sara 1,298, etc.). "
    "Si un mismo negocio tiene reseña en ambas fuentes, se muestra solo la de SQLite. "
    "El frontend pagina de 20 en 20 sin nuevas llamadas al backend.")

sep(doc)

# ── Cold-start wizard ──
h2(doc, "Cold-start Wizard")
bwc(doc, [("Endpoints: ", False), ("POST /users/me/coldstart", True),
    (" (guardar) + ", False), ("GET /users/me/coldstart", True), (" (recuperar)", False)])
body(doc,
    "El wizard de 4 pasos (categorías, ocasión, franja horaria, precio) genera un perfil "
    "que se guarda en localStorage del navegador inmediatamente, y en SQLite si el usuario "
    "está autenticado. Al hacer login desde otro dispositivo, GET /users/me/coldstart "
    "recupera el perfil guardado.")
body(doc,
    "El perfil se usa en businesses.py: _profile_to_content_params() lo convierte en "
    "parámetros para get_content_scores_for_city() de coldstart.py. "
    "El resultado es un {business_id: score} que inject_scores() usa como cold_start_scores "
    "para el camino 3 (TF-IDF) o el blend α (camino 1b).")


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Pre-filtros — antes de que inject_scores() corra")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "El endpoint GET /businesses aplica tres filtros en cascada antes de puntuar nada. "
    "Un negocio que no pasa cualquiera de estos nunca llega a inject_scores().")

simple_table(doc,
    ["Filtro", "Reduce a", "Función / archivo"],
    [
        ("Ciudad",           "~10k negocios (Philadelphia: 10,610)",   "list_businesses() — businesses.py"),
        ("Barrio (opcional)","~628 negocios (University City)",         "list_businesses() — businesses.py"),
        ("is_open_now()",    "Negocios abiertos en este momento",       "business_hours.py → _is_open_at()"),
    ]
)

note(doc,
    "Si tras el filtro de horario quedan menos de 20 negocios, el filtro se omite (protección nocturna). "
    "El 15.4% de negocios sin datos de horario en Yelp siempre pasa (fail open).")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Archivos del scorer — dónde está cada pieza")
# ══════════════════════════════════════════════════════════════════════════════

simple_table(doc,
    ["Archivo", "Responsabilidad"],
    [
        ("hybrid_scorer.py",    "Cadena de prioridades, blend α, fórmula final, whyPicked"),
        ("svdpp_engine.py",     "build_warm_user_vector(), compute_folding_in_vector(), score_all_businesses_for_user()"),
        ("coldstart.py",        "_compute_raw_scores(), get_content_scores_for_city()"),
        ("_state.py",           "Variables en memoria: top_n, explanations, svdpp_model, tfidf, biz_pop_cb"),
        ("_loaders.py",         "Carga al startup: parquets, content model, SVD++"),
        ("contextual_scorer.py","live_ctx_score(tags, hour) → CTX [0,1]"),
        ("business_hours.py",   "is_open_now(business_id, dt) → pre-filtro de horario"),
        ("businesses.py (router)","Coordina pre-filtros, cold_start_scores, llama inject_scores()"),
    ]
)


# ── Save ──────────────────────────────────────────────────────────────────────
out = Path(__file__).parent.parent / "Lantern_Quick_Reference.docx"
doc.save(str(out))
print(f"Generado: {out}")
