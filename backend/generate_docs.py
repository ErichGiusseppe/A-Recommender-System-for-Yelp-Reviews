"""
generate_docs.py — genera la documentación técnica de Lantern en Word (.docx).
Corre desde backend/: python generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ── helpers ──────────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1C, 0x19, 0x17)
    return p

def h3(doc, text):
    return doc.add_heading(text, level=3)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    return p

def inline_code(run):
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)

def bwc(doc, parts):
    """Paragraph mixing normal text and inline code. parts = [(text, is_code)]."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    for text, is_code in parts:
        run = p.add_run(text)
        if is_code:
            inline_code(run)
    return p

def code_block(doc, text):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x57, 0x53, 0x4E)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(7)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5F5F4")
    p._p.get_or_add_pPr().append(shd)
    return p

def callout(doc, text, color="FFF7ED"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    p._p.get_or_add_pPr().append(shd)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xE7, 0xE5, 0xE4)
    run.font.size = Pt(8)

def endpoint_table(doc, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["Endpoint", "Descripción"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1C1917")
        cell._tc.get_or_add_tcPr().append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (ep, desc) in enumerate(rows):
        r = tbl.rows[i + 1].cells
        r[0].text = ep
        r[0].paragraphs[0].runs[0].font.name = "Courier New"
        r[0].paragraphs[0].runs[0].font.size = Pt(9)
        r[1].text = desc
        if i % 2 == 0:
            for cell in r:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FAF6F0")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Title page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lantern — Documentación Técnica de Flujos")
run.bold = True; run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("MINE 4201 · Taller 2 · 2026-1 · Universidad de los Andes")
r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
doc.add_paragraph()

callout(doc,
    "Este documento explica, de principio a fin, cómo funciona cada pantalla y cada flujo de Lantern: "
    "qué ve el usuario, qué endpoint se llama, qué función lo procesa y qué decide el motor de recomendación. "
    "Está escrito para alguien del equipo que conoce sistemas de recomendación pero que no ha tocado este código antes.",
    color="EFF6FF")


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "0. Qué es Lantern y con qué datos trabaja")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Lantern es una app de recomendación de restaurantes y negocios construida sobre el dataset público "
    "de Yelp. La idea es sencilla: el usuario elige una ciudad, y la app le muestra los lugares que "
    "más le van a gustar según su historial de calificaciones, el momento del día y sus preferencias "
    "declaradas. Todo el razonamiento de 'por qué te recomendamos esto' es visible en la interfaz — "
    "cada negocio muestra cuánto vino del historial del usuario, cuánto de sus gustos declarados y "
    "cuánto de que simplemente es popular en el área.")

body(doc,
    "El dataset de Yelp tiene 6,990,280 reseñas, 1,987,897 usuarios y 150,346 negocios en 10 ciudades "
    "de Estados Unidos y Canadá: Philadelphia, Indianapolis, Nashville, Tampa, New Orleans, Tucson, "
    "Reno, Edmonton, Saint Louis y Santa Barbara. Lantern carga todos estos negocios al iniciar el "
    "servidor y los mantiene en memoria para responder sin consultar disco en cada request.")

body(doc,
    "El motor de recomendación combina tres tipos de señal. La señal colaborativa (CF) viene de un "
    "modelo SVD++ entrenado con el dataset Yelp: aprende qué usuarios con gustos similares valoran "
    "los mismos negocios. La señal de contenido (CB) se activa para usuarios sin historial: usa un "
    "modelo TF-IDF que compara las preferencias declaradas del usuario con las categorías de cada "
    "negocio. La señal contextual (CTX) ajusta el ranking según la hora del día: un bar de cócteles "
    "sube en el ranking a las 9pm aunque tenga el mismo score SVD++ que a las 10am.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Términos clave — leer esto primero")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Antes de entrar a los flujos, es útil tener claros estos términos que aparecen en todo el código y "
    "en este documento.")

h3(doc, "Usuario warm vs usuario cold")
body(doc,
    "Un usuario warm es alguien cuyo user_id está en el trainset del modelo SVD++. Esto solo aplica a "
    "los 7 usuarios demo del proyecto (ver sección 2), que fueron seleccionados del dataset Yelp antes "
    "del entrenamiento. El modelo tiene un vector latente de 25 dimensiones para cada uno, que "
    "representa sus gustos en el espacio matemático compartido con todos los negocios. Un usuario cold "
    "es cualquiera que se registró en la app: su user_id es un UUID nuevo que el modelo nunca vio, "
    "así que no hay vector para él y el sistema debe inferir sus gustos de otra forma.")

h3(doc, "SVD++ y vector latente")
body(doc,
    "SVD++ (Singular Value Decomposition++) es el algoritmo de filtrado colaborativo que usa Lantern. "
    "Se entrenó durante ~26 minutos en Colab con 5,592,224 reseñas. El resultado es dos matrices: P "
    "con un vector de 25 números por cada usuario (su 'huella de gustos') y Q con un vector de 25 "
    "números por cada negocio (su 'huella de características'). La predicción de qué tanto le va a "
    "gustar el negocio B al usuario A es básicamente el producto punto de sus dos vectores más "
    "algunos sesgos. Lo importante: este modelo solo conoce a los usuarios que estaban en el dataset "
    "de entrenamiento. Usuarios nuevos no tienen vector P.")

h3(doc, "Folding-in")
body(doc,
    "Cuando un usuario warm califica negocios en la app, el sistema puede actualizar su vector P "
    "sin reentrenar el modelo completo. Esto se llama folding-in: se resuelve un sistema de mínimos "
    "cuadrados con las nuevas calificaciones y la matriz Q fija. Es una actualización en tiempo real "
    "que tarda menos de 50ms. Para usuarios cold también funciona, pero al no tener vector P previo "
    "el sistema empieza desde cero y necesita más reseñas para que el resultado sea confiable.")

h3(doc, "CF, CB, CTX, POP y match")
body(doc,
    "Cada negocio en la respuesta del API tiene cinco campos numéricos (todos de 0 a 100): CF es la "
    "contribución del modelo colaborativo, CB la del modelo de contenido, CTX la del contexto horario "
    "y POP la de popularidad. El campo match es el score final visible al usuario, calculado como "
    "60% CF + 25% CTX + 15% POP. CB reemplaza a CF cuando el usuario no tiene historial colaborativo. "
    "whyPicked es un texto generado automáticamente que explica qué señal dominó el score de ese "
    "negocio para ese usuario en ese momento.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Usuarios del sistema: demos, registrados e invitados")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Hay tres tipos de usuario y cada uno tiene un flujo de recomendación distinto.")

body(doc,
    "Los siete usuarios demo (camila, daniel, sara, alex, maria, carlos, sofia) son cuentas especiales "
    "cuyos user_ids corresponden a usuarios reales del dataset Yelp que fueron seleccionados como los "
    "top reviewers de Philadelphia. Se eligieron exactamente siete para poder demostrar personalización "
    "real: cada uno tiene un vector SVD++ diferente entrenado con sus reseñas históricas de Yelp "
    "(Camila tiene 3,048 reseñas, Daniel 1,653, Sara 1,298...). Sus credenciales son simples: "
    "username = nombre en minúscula, password = mismo nombre. Al autenticarse, el backend busca el "
    "username en data/user_profiles_map.json, valida la contraseña y devuelve un JWT con el user_id "
    "real del dataset. Gracias a esto, inject_scores() puede encontrar su vector en el modelo SVD++ "
    "y generar recomendaciones realmente personalizadas.")

body(doc,
    "Los usuarios registrados se crean en la app con POST /auth/register. Su user_id es un UUID "
    "generado en el momento del registro y se almacena en SQLite. Como el modelo SVD++ nunca los "
    "vio, arrancan en modo cold-start: el sistema los trata como alguien sin historial y usa el "
    "wizard de preferencias para personalizar.")

body(doc,
    "Los usuarios invitados no tienen cuenta. El backend los identifica como 'new_visitor' y les "
    "muestra un ranking de popularidad por ciudad. Pueden completar el wizard de preferencias para "
    "mejorar sus recomendaciones sin necesidad de registrarse.")

callout(doc,
    "Importante: para los endpoints críticos como calificar negocios o guardar preferencias, el "
    "backend exige autenticación. Si un invitado intenta calificar, recibe HTTP 401. Para Discovery "
    "y Search, la autenticación es opcional: si no hay token, el sistema funciona en modo invitado.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Primera apertura de la app")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "La primera vez que alguien abre Lantern, no hay ciudad seleccionada ni perfil guardado. El "
    "sistema muestra un wizard de bienvenida que le pide al usuario elegir una ciudad y un barrio. "
    "Esto se guarda en localStorage del navegador con las claves lantern_city y lantern_neighborhood. "
    "A partir de ese momento, todas las llamadas al backend incluyen city y neighborhood como "
    "parámetros, y el usuario ve solo negocios de esa área.")

body(doc,
    "Después de elegir la ciudad, el wizard de preferencias puede aparecer (también al abrir nuevas "
    "pestañas o en la primera visita de cada sesión). Este wizard de 4 pasos recoge categorías de "
    "interés, la ocasión habitual (viajando, salida local, cita, algo rápido), la franja horaria "
    "preferida y el rango de precio. El perfil resultante se guarda localmente y, si el usuario está "
    "autenticado, también en el backend via POST /users/me/coldstart. Con este perfil, el sistema "
    "puede personalizar aunque no haya historial de calificaciones.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Página Discovery — Tonight in [Ciudad]")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Discovery es la pantalla principal. El titular 'Tonight in Philadelphia' o 'Morning in "
    "Indianapolis' cambia según la hora del día y la ciudad seleccionada. Muestra cinco secciones: "
    "Top Picks, Because You Liked, Trending Nearby, Hidden Gems y Best Value. Las cinco se "
    "construyen sobre los mismos datos: una lista de hasta 50 negocios que el backend calcula y "
    "ordena específicamente para este usuario en este momento.")

h2(doc, "4.1 La llamada al backend — cómo llegan los 50 negocios")

body(doc,
    "Al cargar Discovery, el frontend hace una única llamada al backend que tarda entre 2 y 50ms "
    "dependiendo de si el modelo SVD++ tiene que calcular scores en vivo:")
code_block(doc, "GET /businesses?city=Philadelphia&neighborhood=University+City&limit=50")

body(doc,
    "El backend (app/routers/businesses.py, función list_businesses()) aplica tres filtros en "
    "cascada antes de puntuar nada. Primero filtra por ciudad, reduciendo de los ~150k negocios "
    "globales a los ~10,610 de Philadelphia. Luego filtra por barrio si está seleccionado: University "
    "City tiene 628 negocios. Finalmente aplica el filtro de horario: consulta si cada negocio está "
    "abierto ahora mismo según sus datos de horario de Yelp. Un restaurante que cierra a las 10pm "
    "no aparece en la lista a las 11pm. Si tras este filtro quedan menos de 20 negocios (por ejemplo "
    "a las 3am), el filtro de horario se omite para evitar dejar la pantalla vacía.")

body(doc,
    "Con los candidatos ya filtrados, el backend llama a inject_scores() del módulo "
    "app/services/recommender/hybrid_scorer.py. Esta función recorre todos los negocios candidatos "
    "y asigna a cada uno los campos match, cf, cb, ctx y pop. Luego los ordena de mayor a menor "
    "match y devuelve los primeros 50. Esos 50 son exactamente los que ve el usuario en Discovery, "
    "distribuidos entre las cinco secciones.")

h2(doc, "4.2 Cómo calcula inject_scores() el match de cada negocio")

body(doc,
    "Dentro de inject_scores() el sistema evalúa al usuario actual y toma el camino que corresponda:")

body(doc,
    "Si el usuario tiene 5 o más reseñas en la app, se activa el folding-in. El sistema toma esas "
    "reseñas de SQLite y resuelve un sistema de mínimos cuadrados para actualizar el vector de "
    "preferencias del usuario manteniendo el modelo fijo. Con ese vector recalculado puntúa todos "
    "los negocios candidatos. Este camino tarda ~50ms pero da la señal más actualizada.")

body(doc,
    "Si el usuario es uno de los 7 demos o cualquier usuario que estaba en el trainset de SVD++ "
    "pero no tiene reseñas nuevas en la app, el sistema recupera su vector directamente del modelo "
    "(ya fue calculado durante el entrenamiento) y puntúa todos los negocios con una multiplicación "
    "matricial en NumPy. Este camino tarda ~2ms y es el más común para los demos.")

body(doc,
    "Si el usuario tiene entre 1 y 4 reseñas en la app, el sistema combina las dos señales: la "
    "colaborativa (que ya empieza a formarse con pocas reseñas) y la de contenido del wizard. La "
    "mezcla es progresiva: con 1 reseña el resultado es 20% colaborativa y 80% contenido; con 4 "
    "reseñas es 80% colaborativa y 20% contenido. Esto aplica igual para usuarios demos y para "
    "nuevos registrados, garantizando un comportamiento consistente.")

body(doc,
    "Si el usuario no tiene historial ni reseñas en la app pero sí completó el wizard, se usa el "
    "perfil de preferencias para construir un query TF-IDF y calcular similitud coseno contra los "
    "vectores de categorías de los negocios. El negocio que más coincide con las categorías "
    "preferidas del usuario, ajustado por popularidad, sube en el ranking.")

body(doc,
    "Si no hay nada — usuario invitado sin wizard — se usa un ranking de popularidad pre-calculado "
    "para la ciudad. Es el mismo para todos los invitados de esa ciudad.")

body(doc, "En todos los casos el score final sigue la misma fórmula:")
code_block(doc, "match = ⌊100 · (0.60 · CF  +  0.25 · CTX  +  0.15 · POP)⌋")

body(doc,
    "CF es la señal colaborativa (o de contenido si no hay historial). CTX es el boost contextual "
    "calculado en vivo con la hora actual del request: a las 8am, los cafés y restaurantes de "
    "desayuno reciben un boost de hasta 90 puntos; a las 10pm, los bares y pizzerías. POP es la "
    "popularidad normalizada del negocio basada en su número de reseñas. Los pesos (60/25/15) "
    "están definidos en app/config.py y se pueden ajustar desde un solo lugar.")

sep(doc)

h2(doc, "4.3 Sección: Top Picks")

body(doc,
    "Top Picks muestra los 4 negocios con mayor match de la lista. El titular 'Tuned to your "
    "answers' o 'Personalized · just now' indica que la selección es personalizada para este usuario.")

body(doc,
    "Hay un caso especial importante: si el usuario completó el wizard de preferencias y el sistema "
    "detecta que no tiene señal colaborativa real (ningún negocio de la lista tiene CF > 10), entonces "
    "Top Picks muestra los resultados de una llamada adicional al endpoint de cold-start en lugar de "
    "los 4 primeros de la lista general. El frontend detecta si el usuario es 'warm' revisando si "
    "algún negocio de la lista tiene CF > 10 — si es así, siempre usa los 4 primeros del ranking SVD++.")

body(doc,
    "Cuando se necesitan picks de cold-start, el frontend llama a:")
code_block(doc, "GET /recommendations/cold-start?categories=Italian,Coffee&price=2&stars=0.8&city=Philadelphia")

bwc(doc, [
    ("El router en ", False), ("app/routers/recommendations.py", True),
    (" convierte esos parámetros en un query TF-IDF y devuelve los negocios con mayor afinidad al "
     "perfil del usuario. El motivo de usar cold-start aquí en lugar del ranking general es que para "
     "usuarios sin historial colaborativo, el ranking general estaría dominado por popularidad y "
     "contexto horario, ignorando las preferencias que el usuario se tomó el tiempo de declarar.", False)
])

callout(doc,
    "Si el perfil del wizard de una sesión anterior persiste en el navegador pero el usuario actual "
    "tiene señal SVD++ real (CF > 10), el sistema ignora el wizard y usa SVD++ para Top Picks. Esto "
    "previene que el perfil de un usuario anterior contamine las recomendaciones del siguiente.")

sep(doc)

h2(doc, "4.4 Sección: Because You Liked [Negocio]")

body(doc,
    "Esta sección dice, por ejemplo, 'Because you liked Zahav' y muestra tres negocios similares. "
    "El nombre del negocio ancla es el que tiene el mayor score CF en la lista, dentro de una "
    "categoría específica como 'Mediterranean' o 'Coffee & Tea'. No hace ninguna llamada adicional "
    "al backend — opera sobre los mismos 50 negocios ya cargados.")

body(doc,
    "El proceso de selección en el frontend funciona así: filtra los negocios con CF > 50 (señal "
    "colaborativa fuerte), descarta categorías genéricas como Restaurants, Food o Nightlife que no "
    "dicen nada específico, y agrupa los restantes por categoría. De las tres categorías con más "
    "negocios representados, elige una al azar — esto hace que la sección varíe entre sesiones. "
    "Por ejemplo, si Camila tiene señal fuerte para Mediterranean (5 negocios) y para Coffee & Tea "
    "(3 negocios), la sección puede mostrar cualquiera de las dos.")

body(doc,
    "Para que una categoría sea elegible necesita tener al menos 3 negocios disponibles, de forma "
    "que siempre se puedan mostrar el ancla más 2 recomendaciones. Si ninguna categoría llega a ese "
    "mínimo, la sección no aparece. El número mostrado en el header ('3 nearby in this lane') es "
    "dinámico y refleja cuántas recomendaciones realmente se encontraron.")

sep(doc)

h2(doc, "4.5 Sección: Trending Nearby")

body(doc,
    "Trending muestra los 3 negocios con mayor match score entre todos los de la lista, con un "
    "filtro adicional: el rating debe ser de al menos 3.5 estrellas. Sin ese filtro, un negocio con "
    "2.5 estrellas pero muy popular podría aparecer como 'trending', lo cual sería contradictorio. "
    "Como la lista ya viene ordenada por match, no hay lógica adicional de ranking — simplemente "
    "se toman los 3 primeros que cumplan el rating mínimo.")

sep(doc)

h2(doc, "4.6 Sección: Hidden Gems Nearby")

body(doc,
    "Hidden Gems busca lugares con calificación alta (≥ 4.5 estrellas) pero pocas reseñas — el "
    "indicador de que son buenos pero aún no muy conocidos. Se ordenan de menor a mayor número de "
    "reseñas y se toman los primeros 3. Esta sección no usa personalización: el resultado es el "
    "mismo para todos los usuarios en el mismo barrio. Puede incluir cualquier tipo de negocio — "
    "una peluquería con 5 estrellas y 8 reseñas puede aparecer junto a un restaurante con 4.5 y "
    "12 reseñas. El objetivo es mostrar lo desconocido y excelente, sin importar la categoría.")

sep(doc)

h2(doc, "4.7 Sección: Best Value")

body(doc,
    "Best Value muestra los 3 negocios con mejor calificación entre los económicos (precio $ o $$). "
    "La condición adicional es que el dato de precio debe venir del dataset Yelp real, no ser un "
    "valor por defecto. En el parquet de negocios, cada negocio tiene un campo price_known que es "
    "true solo si Yelp tenía datos de precio para ese negocio. Aproximadamente el 30% de los negocios "
    "en ciudades como Indianapolis no tienen precio en Yelp y reciben '$$' como etiqueta por defecto "
    "— esos se excluyen de Best Value para no mostrar negocios potencialmente caros como 'económicos'.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Página Search")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Search permite explorar negocios con texto libre y filtros combinables. El usuario puede buscar "
    "por nombre o descripción, filtrar por una o varias categorías a la vez, por rango de precio y "
    "por ciudad. Cada cambio en cualquier filtro dispara una búsqueda automáticamente con 300ms de "
    "retraso para no saturar el backend mientras el usuario escribe.")

bwc(doc, [
    ("La llamada es ", False), ("GET /search", True),
    (" con los parámetros que apliquen. El backend en ", False),
    ("app/routers/search.py", True),
    (" primero filtra el catálogo completo por texto (busca en nombre, categoría, barrio y etiquetas), "
     "luego aplica los filtros de categoría y precio, y finalmente llama a inject_scores() para "
     "ordenar los resultados por relevancia personal del usuario — exactamente igual que en Discovery. "
     "Esto significa que si Camila busca 'pizza', los resultados aparecen ordenados según su "
     "historial SVD++, no simplemente por popularidad.", False)
])

body(doc,
    "La búsqueda de texto no es semántica — es una búsqueda de texto plano que verifica si la "
    "query aparece en el nombre del negocio, su categoría, el barrio o sus etiquetas normalizadas. "
    "Buscar 'cozy Italian' no funciona como en un buscador moderno; buscar 'Italian' sí filtra "
    "correctamente porque 'Italian' aparece en la categoría de muchos negocios. El filtro de "
    "categoría es un multi-select: seleccionar 'Italian' y 'Mediterranean' muestra ambos.")

body(doc,
    "El mapa que aparece a la derecha de los resultados se centra automáticamente en el barrio o "
    "ciudad seleccionada — no está fijo en Philadelphia como podría sugerir el nombre del componente "
    "PhillyMap en el código.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Página Detail — Detalle de un negocio")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Al entrar al detalle de un negocio, el usuario ve la galería de fotos, las reseñas reales de "
    "Yelp, los horarios y la ExplanationCard — la tarjeta que explica por qué se recomendó ese "
    "lugar específicamente para ese usuario. El frontend hace dos llamadas en paralelo al cargar "
    "la página.")

bwc(doc, [
    ("La primera, ", False), ("GET /businesses/{id}", True),
    (", retorna el negocio con sus scores calculados en tiempo real. El backend corre inject_scores() "
     "igual que para la lista, garantizando que los valores CF/CB/CTX/POP que ve el usuario en el "
     "detail sean los mismos que tuvo en la lista principal.", False)
])

bwc(doc, [
    ("La segunda, ", False), ("GET /explanations/{id}", True),
    (", busca la explicación de por qué ese negocio apareció. Primero busca una explicación "
     "personalizada para el par usuario-negocio específico (precalculada al generar los parquets). "
     "Si no existe, usa la explicación del usuario invitado para ese negocio. Si tampoco hay, "
     "usa los scores calculados en tiempo real del primer request.", False)
])

body(doc,
    "La ExplanationCard desglosa el match en sus componentes: qué porcentaje vino del historial "
    "colaborativo, de las preferencias declaradas, del contexto horario y de popularidad. El texto "
    "que acompaña cada señal es dinámico — dice 'Usuarios con historial similar valoraron este "
    "lugar altamente' si CF domina, 'Perfecto para este momento del día' si CTX domina, y así. "
    "El campo whyPicked de las cards de lista también es dinámico y refleja la señal dominante "
    "para ese negocio específico: puede decir 'Picked mostly your taste history' para un negocio "
    "con CF alto, o 'Picked perfect for this time of day' para uno que sube por el boost nocturno.")

h2(doc, "6.1 Calificar un negocio")

body(doc,
    "Desde el detail el usuario puede dejar una calificación de 1 a 5 estrellas. Al confirmar, "
    "el frontend envía POST /reviews con el business_id, el número de estrellas y el texto opcional. "
    "El backend guarda la reseña en SQLite con INSERT OR REPLACE — si ya había una reseña del mismo "
    "usuario para ese negocio, simplemente la actualiza con la nueva calificación.")

body(doc,
    "El efecto en las recomendaciones es inmediato desde el siguiente request. Cuando el usuario "
    "vuelve a Discovery o Search, el backend lee todas sus reseñas de SQLite y activa el mecanismo "
    "apropiado según cuántas tenga: con 1 a 4 reseñas usa el blend progresivo (más peso al wizard, "
    "menos al modelo), y con 5 o más activa el folding-in puro que recalcula su vector de "
    "preferencias en tiempo real. Los negocios ya calificados desaparecen automáticamente de "
    "todas las recomendaciones.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Página Profile (You)")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "El perfil muestra los datos del usuario autenticado: su avatar con iniciales, estadísticas "
    "de uso, una visualización de sus gustos ('Your taste this season') y dos tabs con sus "
    "lugares y reseñas.")

body(doc,
    "Las estadísticas del header se calculan en tiempo real desde la base de datos: el número "
    "de reseñas escritas y el rating promedio dado vienen de una consulta SQL directa a la tabla "
    "reviews de SQLite. No son valores hardcodeados.")

body(doc,
    "Las barras de 'Your taste this season' (Italian 28%, Wine bars 22%, etc.) representan la "
    "distribución de categorías en las top-30 recomendaciones del usuario. El backend calcula "
    "esto en users.py tomando las recomendaciones del parquet SVD++ del usuario, buscando la "
    "categoría de cada negocio recomendado y contando cuántos caen en cada una. Esto significa "
    "que para los usuarios demo refleja sus gustos reales del dataset Yelp; para usuarios nuevos "
    "refleja el perfil del wizard mientras van acumulando reseñas.")

h2(doc, "7.1 Tab Saved")

bwc(doc, [
    ("Al cargar el perfil, ", False), ("GET /users/me", True),
    (" retorna entre otros datos la lista de los primeros 12 business_ids de las recomendaciones "
     "del usuario (el parquet pre-computado). El frontend luego fetcha el detalle de cada uno "
     "individualmente para obtener imagen, rating y precio. Estos son los 'lugares guardados' del "
     "usuario — en realidad son sus mejores recomendaciones del sistema, no lugares que el usuario "
     "marcó explícitamente como favoritos.", False)
])

h2(doc, "7.2 Tab Reviews")

bwc(doc, [
    ("El tab Reviews llama a ", False), ("GET /reviews/me", True),
    (", que combina dos fuentes de datos. La primera son las reseñas escritas en la app, guardadas "
     "en SQLite — estas tienen prioridad. La segunda son las reseñas históricas del usuario en el "
     "dataset Yelp, cargadas desde demo_user_reviews.parquet (un archivo pre-generado que extrae "
     "las reseñas de los 7 usuarios demo: Camila tiene 3,048, Daniel 1,653, Sara 1,298, etc.). "
     "Esto permite que cuando un usuario demo entra a Reviews, vea su historial real de Yelp en "
     "lugar de una pantalla vacía. Si el mismo negocio tiene reseña tanto en SQLite como en el "
     "parquet, se muestra solo la de SQLite (más reciente).", False)
])

body(doc,
    "La lista se muestra paginada de 20 en 20 con un botón 'Load more'. Todos los datos llegan "
    "en una sola llamada al backend — la paginación ocurre solo en el frontend mostrando "
    "progresivamente más elementos de la misma lista.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Flujo Cold-Start — del wizard a las recomendaciones")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "El cold-start resuelve un problema fundamental: cómo personalizar para alguien que el modelo "
    "nunca vio. En Lantern, esto aplica a todos los usuarios nuevos que se registran en la app "
    "(su UUID no está en el modelo SVD++) y a los usuarios invitados.")

body(doc,
    "El wizard de preferencias tiene 4 pasos. El usuario elige hasta 3 categorías de interés "
    "(Italian, Seafood, Coffee...), la ocasión habitual (viajando, salida local, cita romántica, "
    "algo rápido), la franja horaria preferida (mañana, almuerzo, tarde, cena, noche) y el rango "
    "de precio ($, $$, $$$). Al confirmar, el perfil se guarda inmediatamente en el almacenamiento "
    "local del navegador para que esté disponible sin conexión, y si el usuario está autenticado "
    "también se persiste en el backend via POST /users/me/coldstart para recuperarlo en otros "
    "dispositivos.")

body(doc,
    "Cuando el backend necesita recomendaciones para un usuario cold-start, convierte el perfil "
    "del wizard en parámetros para el modelo TF-IDF. Por ejemplo, si el usuario eligió 'dinner' "
    "como franja horaria y 'local' como ocasión, el sistema construye automáticamente un string "
    "de categorías que incluye 'Restaurants, Italian, Steakhouses' (las categorías típicas de "
    "cena) junto con las categorías que el usuario eligió explícitamente. La ocasión 'local' "
    "implica un umbral de calidad de 0.75 sobre 5 estrellas, mientras que 'traveling' sube ese "
    "umbral a 0.92 porque un turista quiere solo lo mejor.")

body(doc,
    "Con ese query, el modelo TF-IDF calcula la similitud coseno entre el perfil del usuario y "
    "el vector de categorías de cada negocio. El score resultante se combina con popularidad "
    "(75% similitud + 25% popularidad), garantizando que incluso para preferencias muy específicas "
    "aparezcan negocios conocidos y no solo lugares desconocidos que coincidan por azar con las "
    "etiquetas.")

body(doc,
    "La transición de cold-start a personalización real es gradual. Con cada reseña que el usuario "
    "deja en la app, el sistema mezcla progresivamente más señal colaborativa y menos de contenido. "
    "Con 1 reseña: 80% wizard, 20% colaborativo. Con 3 reseñas: 40% wizard, 60% colaborativo. "
    "Con 5 reseñas o más: el wizard se ignora completamente y el folding-in toma el control total. "
    "Esta transición funciona igual para cualquier usuario, sea demo o recién registrado.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Pre-filtro de horarios — negocios abiertos ahora")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Antes de calcular ningún score, el endpoint /businesses filtra los negocios que están "
    "actualmente cerrados. Esto se hace consultando los horarios de apertura reales del dataset "
    "Yelp, que tienen el formato {'Monday': '9:0-21:0', 'Tuesday': '9:0-21:0', ...}. El servicio "
    "business_hours.py carga estos datos al arrancar el servidor desde business_hours.parquet, "
    "un archivo generado una sola vez por generate_hours.py que tarda apenas 0.6 segundos en "
    "procesar los 150,346 negocios del dataset.")

body(doc,
    "El parseo de los horarios de Yelp requiere manejar algunos casos no obvios. Cuando todos "
    "los días de un negocio tienen el valor '0:0-0:0', significa que está abierto las 24 horas "
    "(Wawa, Walmart y similares usan este formato en el dataset). Cuando solo algunos días tienen "
    "'0:0-0:0' y otros tienen horas reales, significa que el negocio cierra esos días específicos. "
    "Un horario como '22:0-0:0' significa que cierra a medianoche (el '0:0' en el cierre se "
    "interpreta como las 24:00 del mismo día). Si un día de la semana no aparece en el diccionario, "
    "el negocio está cerrado ese día. Los negocios sin ningún dato de horario siempre pasan el "
    "filtro por defecto.")

body(doc,
    "El 84.6% de los negocios tiene datos de horario en Yelp. El 15.4% restante siempre se "
    "incluye en los resultados. Hay una protección adicional: si tras aplicar el filtro quedan "
    "menos de 20 negocios (por ejemplo a las 4am cuando casi todo está cerrado), el filtro se "
    "omite completamente para no dejar la pantalla vacía. Los horarios del dataset Yelp están en "
    "hora local de cada negocio; el servidor usa su hora local, lo que es una aproximación "
    "aceptable para un prototipo académico.")

sep(doc)


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Referencia rápida de endpoints")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "La documentación interactiva completa está en http://localhost:8000/docs (Swagger UI) cuando "
    "el servidor está corriendo. A continuación el resumen de los 19 endpoints activos:")
doc.add_paragraph()

endpoint_table(doc, [
    ("GET /businesses",                "Lista paginada ordenada por match personal. Params: city, neighborhood, limit, offset."),
    ("GET /businesses/{id}",           "Detalle de un negocio con scores calculados en tiempo real para el usuario actual."),
    ("POST /businesses",               "Crear negocio (auth requerida). Se guarda en SQLite y se agrega al catálogo en memoria."),
    ("GET /search",                    "Búsqueda filtrada por texto, categorías y precio. Params: q, category[], price[], city, limit."),
    ("GET /categories",                "Top 20 categorías por frecuencia de negocios. Pre-computado al startup."),
    ("GET /cities",                    "Lista de ciudades disponibles en el dataset. Pre-computada al startup."),
    ("GET /recommendations",           "Top-N recomendaciones del parquet pre-computado para el usuario autenticado."),
    ("GET /recommendations/cold-start","Top-N por TF-IDF de contenido. Params: categories, price, stars, city, limit."),
    ("GET /explanations/{id}",         "Breakdown CF/CB/CTX/POP para el par usuario autenticado + negocio."),
    ("POST /auth/login",               "Login. Verifica credenciales y devuelve JWT de 24h + datos del usuario."),
    ("POST /auth/register",            "Registro. Crea cuenta en SQLite y devuelve JWT."),
    ("GET /users/me",                  "Perfil completo: nombre, stats reales de SQLite, saved_ids, season_taste."),
    ("POST /users/me/taste",           "Guardar perfil de gustos (endpoint activo, la persistencia está pendiente de implementar)."),
    ("POST /users/me/coldstart",       "Guardar perfil del wizard de preferencias en SQLite."),
    ("GET /users/me/coldstart",        "Recuperar perfil del wizard desde SQLite (null si no existe)."),
    ("GET /users/list",                "Lista de las 7 cuentas demo disponibles para login rápido."),
    ("POST /reviews",                  "Calificar un negocio (auth requerida). Activa blend/folding-in en el siguiente request."),
    ("GET /reviews/me",                "Reseñas del usuario: app (SQLite) + historial Yelp (parquet, solo demos). Paginable."),
    ("GET /health",                    "Estado del servidor: versión del modelo cargado y timestamp de arranque."),
])


# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "11. Archivos clave — dónde tocar qué")
# ══════════════════════════════════════════════════════════════════════════════

body(doc,
    "Si necesitas modificar algo del sistema, estos son los archivos donde vive cada responsabilidad:")

h3(doc, "Configuración global")
bwc(doc, [("app/config.py", True),
    (" — cualquier constante del sistema vive aquí: los pesos del score final (W_CF, W_CTX, W_POP), "
     "paths a los archivos de datos, configuración de JWT y el umbral de reseñas para activar el "
     "folding-in (MIN_RATINGS_FOR_FOLDING_IN = 5). Cambiar un peso aquí afecta todo el sistema.", False)])
bwc(doc, [("app/main.py", True),
    (" — punto de entrada de FastAPI. Registra todos los routers, configura CORS para aceptar "
     "llamadas del frontend, y llama a startup() de cada servicio al arrancar el servidor.", False)])

h3(doc, "Routers — los que reciben los HTTP requests")
bwc(doc, [("app/routers/businesses.py", True),
    (" — el más complejo. Coordina todos los pre-filtros (ciudad, barrio, horario), carga los "
     "ratings del usuario desde SQLite, decide si hay cold-start, llama a inject_scores() y "
     "devuelve la lista ordenada. Si algo falla en Discovery, probablemente está aquí.", False)])
bwc(doc, [("app/routers/search.py", True),
    (" — búsqueda: filtra, puntúa y ordena. Mucho más simple que businesses.py.", False)])
bwc(doc, [("app/routers/reviews.py", True),
    (" — submit de reseñas y listado combinado SQLite + parquet Yelp.", False)])
bwc(doc, [("app/routers/recommendations.py", True),
    (" — endpoints de recomendaciones puras (/recommendations) y explicaciones (/explanations/{id}).", False)])
bwc(doc, [("app/routers/users.py", True),
    (" — perfil del usuario, stats calculados en tiempo real, y endpoints del wizard cold-start.", False)])

h3(doc, "Motor de recomendación")
bwc(doc, [("app/services/recommender/hybrid_scorer.py", True),
    (" — el corazón del sistema. La función score_businesses_for_user() implementa toda la cadena "
     "de prioridades: folding-in, SVD++ warm, blend progresivo α, cold-start TF-IDF, fallback. "
     "También genera el texto dinámico de whyPicked. Si algo falla en la personalización, aquí "
     "se empieza a buscar.", False)])
bwc(doc, [("app/services/recommender/svdpp_engine.py", True),
    (" — la matemática SVD++: recuperar el vector del usuario del modelo entrenado, calcular el "
     "nuevo vector por folding-in con lstsq, y puntuar todos los negocios con multiplicación "
     "matricial en NumPy.", False)])
bwc(doc, [("app/services/recommender/coldstart.py", True),
    (" — el modelo TF-IDF de contenido: construir el query vector desde el perfil del wizard, "
     "calcular similitud coseno y devolver scores por negocio.", False)])
bwc(doc, [("app/services/recommender/_state.py", True),
    (" — variables en memoria compartidas entre todos los módulos del recomendador: los "
     "parquets cargados, el modelo SVD++, el modelo TF-IDF, los lookups de ciudad.", False)])
bwc(doc, [("app/services/recommender/_loaders.py", True),
    (" — funciones que cargan los artefactos desde disco al arrancar: parquets, content model, "
     "SVD++. Cada loader es independiente y falla abierto si el archivo no existe.", False)])

h3(doc, "Catálogo y servicios de datos")
bwc(doc, [("app/services/business_store.py", True),
    (" — mantiene el catálogo completo de negocios en memoria. Carga business_meta.parquet, "
     "las fotos de Yelp, las reseñas de muestra y los excerpts al startup. Expone las funciones "
     "de acceso que usan los routers.", False)])
bwc(doc, [("app/services/business_hours.py", True),
    (" — pre-filtro de horarios: carga los datos de apertura de Yelp y responde si un negocio "
     "está abierto ahora mismo con las reglas de parseo documentadas en la sección 9.", False)])
bwc(doc, [("app/services/contextual_scorer.py", True),
    (" — calcula el boost CTX según la hora actual y las categorías del negocio. Define la tabla "
     "de boost por franja horaria (cafés +90 en la mañana, bares +80 en la noche, etc.).", False)])

h3(doc, "Scripts batch — se corren una sola vez para preparar los datos")
bwc(doc, [("generate_parquets.py", True),
    (" — el más importante. Genera top_n.parquet (recomendaciones pre-computadas por usuario), "
     "explanations.parquet (breakdowns CF/CTX/POP), business_meta.parquet (metadatos de negocios) "
     "y content_model.joblib (el modelo TF-IDF). Tarda ~5-10 minutos.", False)])
bwc(doc, [("generate_reviews.py", True),
    (" — extrae las top-10 reseñas por negocio del dataset Yelp (por votos útiles) y las guarda "
     "en reviews_sample.parquet. Tarda ~2-4 minutos.", False)])
bwc(doc, [("generate_hours.py", True),
    (" — extrae los horarios de apertura de todos los negocios del JSON de Yelp y los guarda en "
     "business_hours.parquet. Tarda ~0.6 segundos.", False)])


# ── Save ──────────────────────────────────────────────────────────────────────
out = Path(__file__).parent.parent / "Lantern_Documentacion_Tecnica.docx"
doc.save(str(out))
print(f"Documento generado: {out}")
